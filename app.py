from flask import Flask, render_template, request, redirect, url_for
from model import ATSOptimizer
import mysql.connector
from docx import Document
from PyPDF2 import PdfReader
import tempfile
import logging
import os
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.logger.setLevel(logging.DEBUG)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'doc', 'docx'}  # Extensões permitidas
app.config['MAX_CONTENT_LENGTH'] = 2 * 1024 * 1024  # Limite de 2MB

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])


optimizer = ATSOptimizer()
optimizer.load_and_train()

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

# Configuração do MySQL
db_config = {
    "host": "localhost",
    "user": "root",
    "password": "",
    "database": "recruitment_ai"
}


def extract_text_from_pdf(pdf_path):
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""  # Evita None
        return text
    except Exception as e:
        print(f"Erro ao ler PDF: {str(e)}")
        return None

def extract_text_from_word(filepath):
    try:
        doc = Document(filepath)
        full_text = []
        special_chars = {"": "->", "•": "-", "–": "-"}

        for para in doc.paragraphs:
            text = para.text.strip()
            for char, replacement in special_chars.items():
                text = text.replace(char, replacement)
            if text:
                full_text.append(text)

        # Extrai texto de tabelas
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                full_text.append(' | '.join(row_text))

        # Extrai texto de cabeçalhos e rodapés
        for section in doc.sections:
            for header in section.header.paragraphs:
                if header.text.strip():
                    full_text.append(header.text.strip())
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    full_text.append(footer.text.strip())

        # Extrai texto de text boxes (formas flutuantes)
        for shape in doc.inline_shapes:
            # Verifica se é um objeto de texto
            if shape.type == 3:  # 3 = Tipo TEXT_BOX
                for paragraph in shape.text_frame.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text.strip())

        return '\n'.join(filter(None, full_text)) if full_text else None

    except Exception as e:
        print(f"Erro ao ler Word: {str(e)}")
        return None

@app.route('/upload_page')
def upload_page():
    db = mysql.connector.connect(**db_config)
    cursor = db.cursor()
    cursor.execute("SELECT title FROM jobs")
    jobs = [job[0] for job in cursor.fetchall()]
    cursor.close()
    db.close()
    return render_template('upload.html', jobs=jobs)

@app.route('/upload', methods=['POST'])
def upload_resume():
    # Verifica qual arquivo foi enviado
    file = None
    file_key = None
    
    # Verifica primeiro se algum arquivo foi enviado
    if not request.files:
        return render_template('optimize.html', error="Nenhum arquivo foi selecionado")
    
    # Procura por qualquer arquivo enviado (PDF ou Word)
    for key in request.files:
        if key in ['pdf_file', 'word_file'] and request.files[key].filename != '':
            file = request.files[key]
            file_key = key
            break
    
    if not file:
        return render_template('optimize.html', error="Nenhum arquivo válido foi selecionado")
    
    job_title = request.form.get('job_title')
    job_description = request.form.get('job_description')
    
    # Verifica se os campos foram preenchidos
    if not job_title or not job_description:
        return render_template('optimize.html', error="Por favor, preencha todos os campos")
    
    try:
        # Salva o arquivo temporariamente
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Extrai o texto conforme o tipo de arquivo
        file_ext = filename.lower().split('.')[-1]
        if file_ext == 'pdf':
            resume_text = extract_text_from_pdf(filepath)
        elif file_ext in ['doc', 'docx']:
            resume_text = extract_text_from_word(filepath)
        else:
            os.remove(filepath)
            return render_template('optimize.html', 
                                error="Tipo de arquivo inválido")
        
        # Validação do texto extraído
        if not resume_text or len(resume_text.strip()) < 50:
            os.remove(filepath)
            return render_template('optimize.html', 
                                error="Não foi possível extrair texto válido do arquivo")

        # Processamento da análise
        result = optimizer.analyze_with_custom_job(job_title, job_description, resume_text)
        
        # Limpeza e salvamento no banco
        os.remove(filepath)
        db = mysql.connector.connect(**db_config)
        cursor = db.cursor()
        cursor.execute(
            "INSERT INTO jobs (title, description, required_skills) VALUES (%s, %s, %s)",
            (job_title, job_description, job_description)
        )
        db.commit()
        cursor.close()
        db.close()
        
        # Retreinamento condicional
        optimizer.check_and_retrain()
        
        return render_template('optimize.html', result=result)
        
    except Exception as e:
        # Limpeza em caso de erro
        if 'filepath' in locals() and os.path.exists(filepath):
            os.remove(filepath)
        app.logger.error(f"Erro no upload: {str(e)}")
        return render_template('optimize.html', 
                            error=f"Erro ao processar: {str(e)}")  

@app.route('/')
def index():
    db = mysql.connector.connect(**db_config)
    cursor = db.cursor()
    cursor.execute("SELECT * FROM jobs")
    jobs = cursor.fetchall()
    cursor.close()
    db.close()
    return render_template('index.html', jobs=jobs)

@app.route('/candidate/<int:candidate_id>')
def candidate(candidate_id):
    db = mysql.connector.connect(**db_config)
    cursor = db.cursor()
    cursor.execute("SELECT * FROM candidates WHERE id = %s", (candidate_id,))
    candidate = cursor.fetchone()
    cursor.close()
    db.close()
    return render_template('candidate.html', candidate=candidate)

@app.route('/optimize', methods=['POST'])
def optimize():
    try:
        candidate_id = int(request.form['candidate_id'])
        target_job = request.form['target_job']
        
        optimizer.load_pretrained_model()  # Carrega o modelo treinado
        result = optimizer.optimize_resume(candidate_id, target_job)
        
        return render_template('optimize.html', result=result)
    except Exception as e:
        return render_template('optimize.html', error=str(e))

if __name__ == '__main__':
    app.run(debug=True)